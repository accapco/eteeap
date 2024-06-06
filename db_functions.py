from app import app, db
from models import *
from sqlalchemy import func
from uuid import uuid4
import os
import sys
import json
from datetime import date, datetime
from pathlib import Path
import time

def _calculate_age(birthyear, birthmonth, birthday):
    today = date.today()
    birth_date = date(year=birthyear, month=birthmonth, day=birthday)
    age = today.year - birthyear
    if (today.month, today.day) < (birthmonth, birthday):
        age -= 1
    return age

def get_system_settings():
    fp = "./cfg.json"
    if os.path.exists(fp):
        with open(fp, 'r') as file:
            cfg = json.loads(file.read())
    else:
        cfg = {
            'academic_year': date.today().year, 'semester': '1st'}
        with open(fp, 'w') as file:
            file.write(json.dumps(cfg))
    return cfg

def update_system_settings(data):
    try:
        fp = "./cfg.json"
        cfg = {
            'academic_year': data['academic_year'],
            'semester': data['semester'],
            'cos_dept_head': data['cos_dept_head'],
            'cit_dept_head': data['cit_dept_head'],
            'cie_dept_head': data['cie_dept_head']
        }
        with open(fp, 'w') as file:
            file.write(json.dumps(cfg))
        return True, "Updated system settings." 
    except Exception as e:
        return False, "Unable to update system settings." 

def get_user(user_id):
    return User.query.filter_by(id=user_id).first().__dict__

def get_users(user_type):
    return [u.__dict__ for u in User.query.filter_by(user_type=user_type).all()]

def get_account_details(user_id):
    user = User.query.filter_by(id=user_id).first().__dict__
    if user['user_type'] == 'student':
        student = Student.query.filter_by(user=user_id).first().__dict__
        user['tup_id'] = student['tup_id']
        user['type_of_student'] = student['type_of_student']
        user['foreign_address'] = student['foreign_address']
    if user['user_type'] == 'instructor':
        instructor = Instructor.query.filter_by(user=user_id).first().__dict__
        college = College.query.filter_by(id=instructor['college']).first().abbr
        user['college'] = college
    return user

def profile(user_id):
    user = User.query.filter_by(id=user_id).first().__dict__
    if user['user_type'] == 'student':
        student = Student.query.filter_by(user=user_id).first().__dict__
        user['tup_id'] = student['tup_id']
        user['type_of_student'] = student['type_of_student']
        user['foreign_address'] = student['foreign_address']
    if user['user_type'] == 'instructor':
        instructor = Instructor.query.filter_by(user=user_id).first().__dict__
        college = College.query.filter_by(id=instructor['college']).first().abbr
        user['college'] = college
    user['educational_background'] = get_educational_backgrounds(user_id)
    user['social_media_accounts'] = get_social_media_accounts(user_id)
    return user

def update_account_details(user_id, data):
    try:
        for key, value in data.items():
            if value == "None":
                data[key] = None
        user = User.query.filter_by(id=user_id).first()
        user.l_name = data['l_name']
        user.f_name = data['f_name']
        user.m_name = data['m_name']
        user.ext_name = data['ext_name']
        user.civil_status = data['civil_status']
        user.citizenship = data['citizenship']
        user.birthmonth = data['birthmonth']
        user.birthday = data['birthday']
        user.birthyear = data['birthyear']
        user.gender = data['gender']
        user.email = data['email']
        user.alternate_email = data['alternate_email']
        user.contact_no = data['contact_no']
        user.local_address = data['local_address']
        if user.user_type == 'student':
            student = Student.query.filter_by(user=user_id).first()
            student.type_of_student = data['type_of_student']
            student.foreign_address = data['foreign_address']
            student.tup_id = data['tup_id']
        elif user.user_type == 'instructor':
            instructor = Instructor.query.filter_by(user=user_id).first()
            college_id = College.query.filter_by(abbr=data['college']).first().id
            instructor.college = college_id
        db.session.commit()
        return True, "Account details have been changed."
    except Exception as e:
        return False, f"An error occured while trying to update account. {e}"

import string
import random

def reset_password(id):
    try:
        user = User.query.filter_by(id=id).first()
        characters = string.ascii_letters
        new_password = ''.join(random.choice(characters) for i in range(6))
        user.reset_password = True
        user.password = new_password
        email = user.email
        db.session.add(user)
        db.session.commit()
        name = f"{user.f_name} {user.l_name}"
        email_message = f"Hi {user.f_name},\n\n\n"
        email_message += "Your request for a password reset has been approved. \n"
        email_message += "Here is your new password:\n\n"
        email_message += f"{new_password}"
        send_email(user.email, "TUP-ETEEAP Request for Password Reset", email_message)
        return True, f"Password for {name} has been reset and sent to {email}."
    except Exception as e:
        return False, f"Unable to reset password. {e}" 

def set_new_password(id, data):
    try:
        user = User.query.filter_by(id=id).first()
        user.password = data['password']
        user.reset_password = False
        db.session.add(user)
        db.session.commit()
        return True, f"Password saved."
    except:
        return False, f"Password was not saved." 

def add_educational_background(uid, data):
    if data['school'] and data['degree'] and data['start_year'] and data['end_year']:
        new_educational_background = EducationalBackground(
            school = data['school'],
            degree = data['degree'],
            start_year = data['start_year'],
            end_year = data['end_year'],
            academic_honors = data['academic_honors'],
            user = uid
        )
        db.session.add(new_educational_background)
        db.session.commit()

def update_educational_background(data):
    educational_background = EducationalBackground.query.filter_by(id=data['id']).first()
    educational_background.school = data['school']
    educational_background.degree = data['degree']
    educational_background.start_year = data['start_year']
    educational_background.end_year = data['end_year']
    educational_background.academic_honors = data['academic_honors']
    db.session.commit()

def get_educational_backgrounds(uid):
    educational_backgrounds = EducationalBackground.query.filter_by(user=uid).all()
    response = []
    for e in educational_backgrounds:
        response.append(e.__dict__)
    return response

def add_social_media_account(uid, data):
    if data['platform'] and data['handle']:
        new_social_media_account = SocialMediaAccount(
            platform = data['platform'],
            handle = data['handle'],
            user = uid
        )
        db.session.add(new_social_media_account)
        db.session.commit()

def update_social_media_account(data):
    social_media_account = SocialMediaAccount.query.filter_by(id=data['id']).first()
    social_media_account.platform = data['platform']
    social_media_account.handle = data['handle']
    db.session.commit()

def get_social_media_accounts(uid):
    social_media_accounts = SocialMediaAccount.query.filter_by(user=uid).all()
    response = []
    for s in social_media_accounts:
        response.append(s.__dict__)
    return response

def complete_ft_login(user_id, data):
    try:
        user = User.query.filter_by(id=user_id).first()
        user.password = data['password']
        if user.user_type == 'student':
            student = Student.query.filter_by(user=user_id).first()
            if student.tup_id != data['tup_id']:
                return False, "Invalid TUP ID."
        update_account_details(user_id, data)
        user.ft_login = False
        db.session.commit()
        return True, "Account Details Saved."
    except Exception as e:
        return False, f"An error occured trying to save account details."

import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import email_auth

def send_email(recipient_email, subject, message):
    sender_email = email_auth.sender_email
    app_password = email_auth.app_password
    # Set up the MIME
    msg = MIMEMultipart()
    msg['From'] = sender_email
    msg['To'] = recipient_email
    msg['Subject'] = subject

    # Add message body
    msg.attach(MIMEText(message, 'plain'))

    # Connect to the SMTP server
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.starttls()
    server.login(sender_email, app_password)

    # Send the email
    text = msg.as_string()
    server.sendmail(sender_email, recipient_email, text)
    server.quit()

def change_password(user_id, data):
    try:
        user = User.query.filter_by(id=user_id).first()
        if user.password == data['current']:
            user.password = data['new']
            db.session.commit()
            return True, "Password changed."
        else:
            return False, "Incorrect password."
    except Exception as e:
        return False, "An error occured when trying to save password."

from operator import itemgetter

def generate_student_report(filters):
    students = get_students()
    data = {
        'total_enrollees': 0,
    }
    if filters['age']:
        data['age'] = {}
    if filters['gender']:
        data['gender'] = {'M': 0, 'F': 0}
    if filters['type_of_student']:
        data['type_of_student'] = {'foreign': 0, 'local': 0}
    if filters['student_course_status']:
        data['course_status'] = {'pending': 0, 'ongoing': 0, 'completed': 0}

    for student in students:
        # get student info
        user = get_user(user_id=student['user'])
        ay = str(student['ay'])
        semester = str(student['semester'])
        program = Program.query.filter_by(abbr=student['program']).first()
        if program == None:
            continue
        else:
            college = College.query.filter_by(id=program.college_id).first()
            if college == None:
                continue
            else:
                college = college.abbr
        # skip if student doesn't match filter
        if filters['ay'] != "All" and ay != filters['ay']:
            continue
        if filters['semester'] != "All" and semester != filters['semester']:
            continue
        if filters['college'] != "All" and college != filters['college']:
            continue
        if filters['program'] != "All" and program != filters['program']:
            continue
        data['total_enrollees'] += 1
        # store data
        if 'age' in data.keys():
            if user['birthyear'] and user['birthmonth'] and user['birthday']:
                age = _calculate_age(user['birthyear'], user['birthmonth'], user['birthday'])
                if age not in data['age']:
                    data['age'][age] = 0
                data['age'][age] += 1
        if 'gender' in data.keys():
            gender = user['gender']
            data['gender'][gender] += 1
        if 'type_of_student' in data.keys():
            if student['type_of_student'] != None:
                type_of_student = student['type_of_student']
                data['type_of_student'][type_of_student] += 1
        if 'course_status' in data.keys():
            subjects = Enrollment.query.filter_by(student=student['id']).all()
            for subject in subjects:
                data['course_status'][subject.status] += 1
    if 'age' in data.keys():
        data['age'] = dict(sorted(data['age'].items()))
    return data

def generate_student_excel_report(filters):
    students = get_students()
    data = {
        'total_enrollees': 0,
        'filters': {
            'academic_year': filters['ay'],
            'semester': filters['semester'],
            'college': filters['college'],
            'program': filters['program'],
            'age_range': f"{filters['lb_age']}-{filters['ub_age']}",
            'gender': filters['gender'],
            'type_of_student': filters['type_of_student']
        },
        'student_list': []
    }

    for student in students:
        user = get_user(user_id=student['user'])
        ay = str(student['ay'])
        semester = str(student['semester'])
        program = Program.query.filter_by(abbr=student['program']).first()
        if user['birthyear'] and user['birthmonth'] and user['birthday']:
            age = _calculate_age(user['birthyear'], user['birthmonth'], user['birthday'])
        else:
            age = "Not set"
        if program == None:
            continue
        else:
            college = College.query.filter_by(id=program.college_id).first()
            if college == None:
                continue
            else:
                college = college.abbr
        if filters['ay'] != "All" and ay != filters['ay']:
            continue
        if filters['semester'] != "All" and semester != filters['semester']:
            continue
        if filters['college'] != "All" and college != filters['college']:
            continue
        if filters['program'] != "All" and program != filters['program']:
            continue
        if not(filters['lb_age'] <= age <= filters['ub_age']):
            continue
        if filters['gender'] != "All" and user['gender'] != filters['gender']:
            continue
        if filters['type_of_student'] != "All" and user['type_of_student'] != filters['type_of_student']:
            continue
        data['total_enrollees'] += 1
        name = f"{student['l_name'].title()}, {student['f_name']} "
        if user['m_name'] not in (None, "None"):
            name += ' ' + user['m_name']
        if user['ext_name'] not in (None, "None"):
            name += ' ' + user['ext_name']
        data['student_list'].append({
            'name': name,
            'college': college,
            'program': program.abbr,
            'age': age,
            'gender': user['gender'],
            'student_type': student['type_of_student']
        })
    data['student_list'] = sorted(data['student_list'], key=itemgetter('age'))
    return data

def generate_faculty_report(filters):
    instructors = get_instructors()
    data = {
        'total_instructors': 0,
        'filters': {
            'academic_year': filters['ay'],
            'semester': filters['semester'],
            'college': filters['college'],
            'programs_taught': filters['programs_taught'],
            'honorarium_status': filters['honorarium_status']
        },
        'faculty_list': []
    }
    instructors = get_instructors()
    for instructor in instructors:
        user = User.query.filter_by(id=instructor['user']).first()
        enrollments = Enrollment.query.filter_by(instructor=instructor['id']).all()
        name = f"{user.l_name.title()}, {user.f_name}"
        if user.m_name not in (None, "None"):
            name += ' ' + user.m_name
        if user.ext_name not in (None, "None"):
            name += ' ' + user.ext_name
        pending = []
        ongoing = []
        completed = []
        honorarium_released = []
        honorarium_onprocess = []
        total_enrollments = 0
        for enrollment in enrollments:
            ay = enrollment.ay
            semester = enrollment.semester
            if filters['ay'] != "All" and ay != filters['ay']:
                continue
            if filters['semester'] != "All" and semester != filters['semester']:
                continue
            student = Student.query.filter_by(id=enrollment.student).first()
            program = Program.query.filter_by(id=student.program).first()
            student = User.query.filter_by(id=student.user).first()
            if program:
                program = program.abbr
            if filters['programs_taught'] != "All" and program != filters['programs_taught']:
                continue
            if filters['honorarium_status'] != "All" and enrollment.honorarium != filters['honorarium_status']:
                continue
            course = Course.query.filter_by(id=enrollment.course).first()
            course_name = course.code
            student_name = f"{student.l_name.title()} {student.f_name}"
            if student.m_name not in (None, "None"):
                student_name += ' ' + student.m_name
            if student.ext_name not in (None, "None"):
                student_name += ' ' + student.ext_name
            if enrollment.status == 'pending':
                pending.append({'course': course_name, 'student': student_name})
            elif enrollment.status == 'ongoing':
                ongoing.append({'course': course_name, 'student': student_name})
            else:
                completed.append({'student': student_name, 'course': course_name})
            if enrollment.honorarium == 'onprocess':
                honorarium_onprocess.append({'course': course_name, 'student': student_name})
            else:
                honorarium_released.append({'course': course_name, 'student': student_name})
            total_enrollments += 1
        if total_enrollments == 0:
            continue
        snapshot = {
            'name': name,
            'program': program
        }
        snapshot['pending'] = {'count': len(pending), 'enrollments': pending}
        snapshot['ongoing'] = {'count': len(ongoing), 'enrollments': ongoing}
        snapshot['completed'] = {'count': len(completed), 'enrollments': completed}
        snapshot['honorarium_released'] = {'count': len(honorarium_released), 'enrollments': honorarium_released}
        snapshot['honorarium_onprocess'] = {'count': len(honorarium_onprocess), 'enrollments': honorarium_onprocess}
        data['faculty_list'].append(snapshot)
        data['total_instructors'] += 1
    return data

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def create_honorarium_pdf(data, report):
    try:
        pdf_filename = f"Honorarium {report.capitalize()} Report.pdf"
        c = canvas.Canvas(pdf_filename, pagesize=letter)
        width, height = letter

        c.setFont("Helvetica-Bold", 12)

        y_position = height - 50
        title_status = "On Process" if report == "onprocess" else "Released"
        c.drawString(width / 2 - 100, y_position, f"Faculty Report - Honorarium {title_status}")
        
        c.setFont("Helvetica", 12)
        filters = data["filters"]
        filter_parts = []

        if filters["academic_year"] != "All":
            filter_parts.append(filters["academic_year"])
        if filters["semester"] != "All":
            if filters["semester"] != "Midyear":
                filter_parts.append(filters["semester"]+" Term")
            else:
                filter_parts.append(filters["semester"])
        if filters["college"] != "All":
            filter_parts.append(filters["college"])
        if filters["programs_taught"] != "All":
            filter_parts.append(filters["programs_taught"])
        if filters["honorarium_status"] != "All":
            filter_parts.append(filters["age_range"])

        filter = ", ".join(filter_parts)
        if filter == "":
            filter = "All"
        y_position -= 20
        c.drawString(50, y_position, f"Filter: {filter}")
        
        y_position -= 20
        for faculty in data["faculty_list"]:
            if faculty[f'honorarium_{report}']['count'] > 0:
                c.setFont("Helvetica-Bold", 12)
                c.drawString(50, y_position, f"Name: {faculty['name']}")
                y_position -= 15
                c.setFont("Helvetica", 12)
                c.drawString(50, y_position, f"Honorarium Count: {faculty[f'honorarium_{report}']['count']}")
                y_position -= 15

                c.drawString(50, y_position, "Enrollments:")
                y_position -= 15
                for enrollment in faculty[f'honorarium_{report}']['enrollments']:
                    c.drawString(70, y_position, f"- {enrollment['student']} ({enrollment['course']})")
                    y_position -= 15

                y_position -= 15

                if y_position <= 50:
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y_position = height - 50

        c.save()
        return True, pdf_filename
    except:
        return False, "Could not create report."

import pandas as pd
from openpyxl import Workbook, load_workbook
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.utils import get_column_letter
from openpyxl.utils.dataframe import dataframe_to_rows 
from openpyxl.styles import Font, Border, Side, Alignment
import os

def create_sheet(data, wb, total_enrollees, sheet_name, Category):
    
    if sheet_name in wb.sheetnames:
        del wb[sheet_name]
    ws = wb.create_sheet(title=sheet_name)
        
    ws.append(["Total Enrollees", total_enrollees])
    ws.append([])
    df = pd.DataFrame(data, columns=[Category, "Count"])
    for r in dataframe_to_rows(df, index=False, header=True):
        ws.append(r)

    chart = BarChart()
    chart.title = sheet_name
    chart.x_axis.title = Category
    chart.y_axis.title = "Count"

    chart.width = 30
    chart.height = 15
    
    temp = Reference(ws, min_col=2, min_row=3, max_col=len(df.columns), max_row=len(df) + 3)
    categories = Reference(ws, min_col=1, min_row=4, max_row=len(df) + 3)
    chart.add_data(temp, titles_from_data=True)
    chart.set_categories(categories)

    ws.add_chart(chart, "E6")

def create_student_excel(data):
    try:
        filters = data["filters"]
        file_name_parts = []

        if filters["academic_year"] != "All":
            file_name_parts.append(filters["academic_year"])
        if filters["semester"] != "All":
            file_name_parts.append(filters["semester"])
        if filters["college"] != "All":
            file_name_parts.append(filters["college"])
        if filters["program"] != "All":
            file_name_parts.append(filters["program"])
        if filters["age_range"] != "All":
            file_name_parts.append(filters["age_range"])
        if filters["gender"] != "All":
            file_name_parts.append(filters["gender"])
        if filters["type_of_student"] != "All":
            file_name_parts.append(filters["type_of_student"])

        file_name = "-".join(file_name_parts) + ".xlsx"
        if file_name == ".xlsx":
            file_name = "All.xlsx"
        excel_file = f'Students - {file_name} - Report.xlsx'

        if os.path.exists(excel_file):
            os.remove(excel_file)

        wb = Workbook()
        wb.remove(wb.active)

        ws = wb.create_sheet(title="Report")
            
        ws.append(["Total Enrollees", data["total_enrollees"]])
        ws.append([])
        ws.append(["Name", "College", "Program", "Age", "Gender", "Student Type"])
        for student in data["student_list"]:
            ws.append([student["name"], student["college"], student["program"], student["age"], student["gender"], student["student_type"]])

        # Auto Fit
        for sheet in wb.worksheets:
            for column in sheet.columns:
                max_length = 7.15
                column = list(column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                sheet.column_dimensions[get_column_letter(column[0].column)].width = adjusted_width

        wb.save(excel_file)
        return True, excel_file
    except Exception as e:
        return False, f"Problem occured while generating report. {e}"

def create_faculty_excel(data):
    try:
        # Lists faculty members and the subjects they handle. Includes details on the number of subjects completed and ongoing for each faculty member.
        filters = data["filters"]
        file_name_parts = []

        if filters["academic_year"] != "All":
            file_name_parts.append(filters["academic_year"])
        if filters["semester"] != "All":
            file_name_parts.append(filters["semester"])
        if filters["college"] != "All":
            file_name_parts.append(filters["college"])
        if filters["programs_taught"] != "All":
            file_name_parts.append(filters["programs_taught"])
        if filters["honorarium_status"] != "All":
            file_name_parts.append(filters["age_range"])

        file_name = "_".join(file_name_parts) + ".xlsx"
        if file_name == ".xlsx":
            file_name = "All.xlsx"
        excel_file = f'Faculty Report {file_name}'

        if os.path.exists(excel_file):
            os.remove(excel_file)

        wb = Workbook()
        wb.remove(wb.active)

        ws = wb.create_sheet(title="Report")
            
        bold_font = Font(bold=True)
        italic_font = Font(italic=True)
        center_center = Alignment(horizontal='center', vertical='center')
        thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

        ws.append(["Total Instructors", data["total_instructors"]])
        ws["A1"].font = bold_font
        ws["B1"].font = bold_font
        ws["A1"].border = thin_border
        ws["B1"].border = thin_border
        ws.append([])

        for faculty in data["faculty_list"]:

            ws.append(["Name", faculty["name"]])
            ws.cell(row=ws.max_row, column=1).font = bold_font
            ws.cell(row=ws.max_row, column=1).border = thin_border
            ws.cell(row=ws.max_row, column=2).font = bold_font
            ws.cell(row=ws.max_row, column=2).border = thin_border
            ws.cell(row=ws.max_row, column=3).font = bold_font
            ws.cell(row=ws.max_row, column=3).border = thin_border
            ws.merge_cells(start_row=ws.max_row, start_column=2, end_row=ws.max_row, end_column=3)

            ws.append(["Program", faculty["program"]])
            ws.cell(row=ws.max_row, column=1).font = bold_font
            ws.cell(row=ws.max_row, column=1).border = thin_border
            ws.cell(row=ws.max_row, column=2).font = bold_font
            ws.cell(row=ws.max_row, column=2).border = thin_border
            ws.merge_cells(start_row=ws.max_row, start_column=2, end_row=ws.max_row, end_column=3)

            enrollment_types = [
                ("Pending Enrollments", faculty["pending"]["count"], faculty["pending"]["enrollments"]),
                ("Ongoing Enrollments", faculty["ongoing"]["count"], faculty["ongoing"]["enrollments"]),
                ("Completed Enrollments", faculty["completed"]["count"], faculty["completed"]["enrollments"]),
                ("Honorarium Released", faculty["honorarium_released"]["count"], faculty["honorarium_released"]["enrollments"]),
                ("Honorarium On Process", faculty["honorarium_onprocess"]["count"], faculty["honorarium_onprocess"]["enrollments"])
            ]

            for enrollment_type, count, enrollments in enrollment_types:
                if count > 0:
                    ws.append([enrollment_type, "Count",count])
                    ws.cell(row=ws.max_row, column=1).font = italic_font
                    ws.cell(row=ws.max_row, column=1).alignment = center_center
                    ws.cell(row=ws.max_row, column=1).border = thin_border
                    ws.cell(row=ws.max_row, column=2).border = thin_border
                    ws.cell(row=ws.max_row, column=3).border = thin_border
                    for enrollment in enrollments:
                        
                            ws.append(["", enrollment["course"], enrollment["student"]])
                            ws.cell(row=ws.max_row, column=1).border = thin_border
                            ws.cell(row=ws.max_row, column=2).border = thin_border
                            ws.cell(row=ws.max_row, column=3).border = thin_border
                    ws.merge_cells(start_row=ws.max_row - len(enrollments), start_column=1, end_row=ws.max_row, end_column=1)
                    ws.append([])               
            ws.append([])
    
        # Auto Fit
        for sheet in wb.worksheets:
            for column in sheet.columns:
                max_length = 7.15
                column = list(column)
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(cell.value)
                    except:
                        pass
                adjusted_width = (max_length + 2)
                sheet.column_dimensions[get_column_letter(column[0].column)].width = adjusted_width

        wb.save(excel_file)
        return True, excel_file
    except Exception as e:
        return False, f"Problem occured while generating report. {e}"

def uid_to_pk(uid, user_type):
    if user_type == 'student':
        obj = Student
    elif user_type == 'instructor':
        obj = Instructor
    elif user_type == 'admin':
        obj = Admin
    else:
        return "Error"
    return obj.query.filter_by(user=uid).first().__dict__['id']

def pk_to_uid(id, user_type):
    if user_type == 'student':
        obj = Student
    elif user_type == 'instructor':
        obj = Instructor
    elif user_type == 'admin':
        obj = Admin
    else:
        return "Error"
    return obj.query.filter_by(id=id).first().__dict__['user']

def add_user(data):
    try:
        for key, value in data.items():
            if value == "None":
                data[key] = None
        new_user = User(
            username=data['username'], 
            f_name=data['f_name'],
            m_name=data['m_name'],
            l_name=data['l_name'],
            ext_name=data['ext_name'],
            email=data['email'],
            gender=data['gender'],
            password=data['password'],
            user_type=data['user_type']
        )
        db.session.add(new_user)
        db.session.commit()
        email_message = f"Hi {data['f_name']}, \n\n\n"
        email_message += f"Here are your TUP-ETEEAP login credentials:\n\n"
        if data['user_type'] == "admin":
            new_admin = Admin(user=new_user.id)
            db.session.add(new_admin)
        elif data['user_type'] == "instructor":
            college_id = College.query.filter_by(abbr=data['college']).first().id
            new_instructor = Instructor(user=new_user.id, college=college_id)
            db.session.add(new_instructor)
        else:
            cfg = get_system_settings()
            new_student = Student(
                user=new_user.id, 
                tup_id=data['tup_id'],
                ay=cfg['academic_year'],
                semester=cfg['semester']
            )
            db.session.add(new_student)
            email_message += f"TUP ID: {data['tup_id']}\n"
        db.session.commit()
        email_message += f"username: {data['username']}\n"
        email_message += f"password: {data['password']}\n\n"
        email_message += f"https://tup-ecadcase-eteeap.onrender.com/"
        try:
            send_email(data['email'], "TUP-ETEEAP Login Credentials", email_message)
            return True, "Successfully added user."
        except:
            return False, "Unable to create account."
    except Exception as e:
        return False, f"Error occured while adding user."

def get_instructor(user_id):
    instructor = Instructor.query.filter_by(user=user_id).first().__dict__
    userinfo = get_user(instructor['user'])
    instructor['f_name'] = userinfo['f_name']
    instructor['l_name'] = userinfo['l_name']
    instructor['m_name'] = userinfo['m_name']
    instructor['username'] = userinfo['username']
    return instructor

def get_instructor_bypk(id):
    instructor = Instructor.query.filter_by(id=id).first().__dict__
    userinfo = get_user(instructor['user'])
    instructor['f_name'] = userinfo['f_name']
    instructor['l_name'] = userinfo['l_name']
    instructor['m_name'] = userinfo['m_name']
    instructor['username'] = userinfo['username']
    return instructor

def get_instructors():
    instructors = []
    for instructor in Instructor.query.all():
        instructor = instructor.__dict__
        userinfo = get_user(instructor['user'])
        instructor['f_name'] = userinfo['f_name']
        instructor['l_name'] = userinfo['l_name']
        instructor['m_name'] = userinfo['m_name']
        instructor['username'] = userinfo['username']
        instructor['college'] = College.query.filter_by(id=instructor['college']).first().abbr
        instructors.append(instructor)
    return instructors

def get_instructor_enrollments(instructor_id):
    enrollments = Enrollment.query.filter_by(instructor=instructor_id).all()
    response = []
    for e in enrollments:
        e = e.__dict__
        e['student'] = get_student_bypk(e['student'])
        e['course'] = get_course(e['course'])
        program = Program.query.filter_by(id=e['student']['program']).first()
        if program:
            e['program'] = program.abbr
            college = College.query.filter_by(id=program.college_id).first()
            if college:
                e['college'] = college.abbr
            else:
                e['college'] = ""
        else:
            e['program'] = ""
        response.append(e)
    return response

def get_academic_years():
    enrollments = Enrollment.query.all()
    academic_years = []
    for enrollment in enrollments:
        if enrollment.ay not in academic_years:
            academic_years.append(enrollment.ay)
    return academic_years

def get_instructor_enrollment(instructor_id, enrollment_id):
    response = Enrollment.query.filter_by(instructor=instructor_id, id=enrollment_id).first().__dict__
    response['student'] = get_student_bypk(response['student'])
    response['course'] = get_course(response['course'])
    return response

def confirm_enrollment(enrollment_id, confirmation):
    enrollment = Enrollment.query.filter_by(id=enrollment_id)
    if confirmation == "accept":
        enrollment = enrollment.first()
        if enrollment.status == 'pending':
            enrollment.status = 'ongoing'
            db.session.commit()
            return "Enrollment accepted."
        else:
            return "Invalid operation"
    elif confirmation == "reject":
        enrollment.delete()
        db.session.commit()
        return "Enrollment rejected."
    else:
        return "Invalid operation"

def get_requirements(enrollment_id):
    requirements = Requirement.query.filter_by(enrollment=enrollment_id).all()
    response = []
    for r in requirements:
        r = r.__dict__
        r['description'] = r['description'].split('\n')
        r['materials'] = get_materials(r['id'])
        r['submissions'] = get_submissions(r['id'])
        response.append(r)
    return response

def get_requirement(requirement_id):
    response = Requirement.query.filter_by(id=requirement_id).first().__dict__
    response['materials'] = get_materials(requirement_id)
    response['description'] = response['description'].split('\n')
    response['submissions'] = get_submissions(requirement_id)
    return response

def add_requirement(enrollment_id, form):
    enrollment = Enrollment.query.filter_by(id=enrollment_id).first()
    new_requirement = Requirement(
        enrollment = enrollment.id,
        title = form.title.data,
        description = form.description.data,
        progress = 'incomplete'
    )
    try:
        db.session.add(new_requirement)
        db.session.commit()
        for file in form.materials.data:
            filename = str(uuid4()) + os.path.splitext(file.filename)[1]
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], 'materials', filename)
            file.save(filepath)
            new_material = RequirementMaterial(filepath=filepath, filename=file.filename, requirement=new_requirement.id)
            db.session.add(new_material)
        db.session.commit()
        return "Requirement posted."
    except Exception as e:
        db.session.rollback()
        return f"An error occured while posting the requirement: {e}"

def get_materials(requirement_id):
    materials = RequirementMaterial.query.filter_by(requirement=requirement_id).all()
    response = []
    for m in materials:
        response.append(m.__dict__)
    return response

def get_submissions(requirement_id):
    submissions = RequirementSubmission.query.filter_by(requirement=requirement_id).all()
    response = []
    for s in submissions:
        response.append(s.__dict__)
    return response

def add_submission(requirement_id, form):
    try:
        for file in form.submissions.data:
            filename = str(uuid4()) + os.path.splitext(file.filename)[1]
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], 'submissions', filename)
            file.save(filepath)
            new_submission = RequirementSubmission(filepath=filepath, filename=file.filename, requirement=requirement_id)
            db.session.add(new_submission)
        db.session.commit()
        return "Requirement posted."
    except Exception as e:
        db.session.rollback()
        return f"An error occured while posting the requirement: {e}"

def turn_in_submission(requirement_id):
    requirement = Requirement.query.filter_by(id=requirement_id).first()
    if requirement.progress == 'incomplete':
        requirement.progress = 'evaluation'
        db.session.commit()
        return "Successfully turned in."
    else:
        return "Invalid operation."
    
def return_submission(requirement_id):
    requirement = Requirement.query.filter_by(id=requirement_id).first()
    if requirement.progress == 'evaluation':
        requirement.progress = 'completed'
        db.session.commit()
        return "Successfully returned."
    else:
        return "Invalid operation."

def get_student(user_id):
    student = Student.query.filter_by(user=user_id).first().__dict__
    userinfo = get_user(student['user'])
    student['f_name'] = userinfo['f_name']
    student['l_name'] = userinfo['l_name']
    student['m_name'] = userinfo['m_name']
    if userinfo['ext_name'] not in [None, "None"]:
        student['ext_name'] = userinfo['ext_name']
    student['username'] = userinfo['username']
    return student

def get_student_bypk(id):
    student = Student.query.filter_by(id=id).first().__dict__
    userinfo = get_user(student['user'])
    student['f_name'] = userinfo['f_name']
    student['l_name'] = userinfo['l_name']
    student['m_name'] = userinfo['m_name']
    if userinfo['ext_name'] not in [None, "None"]:
        student['ext_name'] = userinfo['ext_name']
    student['username'] = userinfo['username']
    return student

def get_students():
    students = []
    for student in Student.query.all():
        student = student.__dict__
        userinfo = get_user(student['user'])
        student['f_name'] = userinfo['f_name']
        student['l_name'] = userinfo['l_name']
        student['m_name'] = userinfo['m_name']
        student['username'] = userinfo['username']
        if student['program']:
            program = Program.query.filter_by(id=student['program']).first()
            student['program'] = program.abbr
            college = College.query.filter_by(id=program.college_id).first().abbr
            student['college'] = college
        students.append(student)
    return students

def get_student_count():
    payment_count = db.session.query(func.count(Student.id)).filter(Student.progress == 'payment').scalar()
    evaluation_count = db.session.query(func.count(Student.id)).filter(Student.progress == 'evaluation').scalar()
    enrollment_count = db.session.query(func.count(Student.id)).filter(Student.progress == 'enrollment').scalar()
    completion_count = db.session.query(func.count(Student.id)).filter(Student.progress == 'completion').scalar()
    return {
        'payment': payment_count,
        'evaluation': evaluation_count,
        'enrollment': enrollment_count,
        'completion': completion_count
    }

def _format_date(time):
    if time:
        return time.strftime("%m/%d/%Y")
    return ""

def get_student_enrollments(student_id):
    enrollments = (
        Enrollment.query.filter_by(student=student_id)
        .join(Instructor)
        .join(Course)
        .all()
    )
    response = []
    for e in enrollments:
        enrollment_data = e.__dict__
        instructor_data = get_instructor_bypk(e.instructor)
        course_data = get_course(e.course)
        enrollment_data['instructor'] = instructor_data
        enrollment_data['course'] = course_data
        enrollment_data['honorarium_released_date'] = _format_date(enrollment_data['honorarium_released_date'])
        response.append(enrollment_data)
    return response

def get_student_enrollment_options(student_id):
    enrolled_classes = get_student_enrollments(student_id)
    all_courses = [u.__dict__ for u in Course.query.all()]
    existing_courses_ids = [c['course']['id'] for c in enrolled_classes]
    available_courses = [course for course in all_courses if course['id'] not in existing_courses_ids]
    response = {
        'student': get_student_bypk(student_id),
        'programs': get_programs(),
        'available_courses': available_courses,
        'instructors': get_instructors(),
        'enrolled': enrolled_classes
    }
    return response

def upload_receipt(user_id, data):
    student = Student.query.filter_by(user=user_id).first()
    student.receipt_filepath = str(uuid4()) + os.path.splitext(data['file'].filename)[1]
    student.progress = 'payment-pending'
    fp = os.path.join(app.config['UPLOAD_FOLDER'], 'receipts', student.receipt_filepath)
    data['file'].save(fp)
    db.session.commit()

def get_receipt_fp(user_id):
    student = Student.query.filter_by(user=user_id).first()
    fp = student.receipt_filepath
    return fp

def accept_receipt(student_id):
    try:
        student = Student.query.filter_by(id=student_id).first()
        student.progress = 'enrollment'
        db.session.commit()
        return True, "Documents have been approved."
    except:
        return False, "Invalid operation."
    
def reject_receipt(student_id):
    try:
        student = Student.query.filter_by(id=student_id).first()
        student.progress = 'payment-rejected'
        db.session.commit()
        return True, "Documents have been rejected."
    except:
        return False, "Invalid operation."

def get_enrollment(id):
    return Enrollment.query.filter_by(id=id).first().__dict__

from itertools import groupby

def get_enrollments_grouped_by_student(instructor_id, filters, get_hfr_and_cp=False):
    enrollments = get_instructor_enrollments(instructor_id)
    response = []
    student_data_list = []
    hfr = 0
    cp = 0
    for e in enrollments:
        if filters['ay'] != "All" and filters['ay'] != e['ay']:
            continue
        if filters['semester'] != "All" and filters['semester'] != e['semester']:
            continue
        if filters['college'] != "All" and filters['college'] != e['college']:
            continue
        if filters['program'] != "All" and filters['program'] != e['program']:
            continue
        if filters['status'] != "All" and filters['status'] != e['status']:
            continue
        student = {
            'name':  f"{e['student']['l_name'].title()}, {e['student']['f_name']} {e['student']['m_name'][0]+'.' if e['student']['m_name'] else ''}",
            'program': e['program'],
            'college': e['college'],
            'tup_id': e['student']['tup_id'],
            'id': e['student']['user'],
            'enrollments': []
        }
        enrollment_data = {
            'course': f"{e['course']['code']}: {e['course']['title']}",
            'status': e['status'],
            'enrollment_id': e['id'],
            'grade': e['grade'],
            'ay': e['ay'],
            'semester': e['semester'],
            'honorarium': e['honorarium'],
            'honorarium_released_date': _format_date(e['honorarium_released_date'])
        }
        student_data_list.append((student, enrollment_data))
        if e['honorarium'] == 'onprocess':
            hfr += 1
        if e['status'] == 'pending':
            cp += 1
    iterator = groupby(student_data_list, lambda x : x[0]) 
    grouping = {}
    for student, enrollments in iterator: 
        grouping = student
        grouping['enrollments'] = [value for key, value in list(enrollments)]
        response.append(grouping)
    if get_hfr_and_cp:
        response = [response, hfr, cp]
    return response

def get_programs_by_college():
    response = []
    for c in College.query.all():
        c = c.__dict__
        c['programs'] = [p.__dict__ for p in Program.query.filter_by(college_id=c['id']).all()]
        response.append(c)
    return response

def get_programs():
    return [p.__dict__ for p in Program.query.all()]

def get_program(program_id):
    program = Program.query.filter_by(id=program_id).first()
    if program:
        return program.__dict__
    else:
        return {
            'id': None,
            'name': 'Not yet enrolled to any program.',
            'abbr': 'Not yet enrolled to any program.',
            'courses': [],
            'college_id': None
        }

def add_course(course):
    try:
        new_course = Course(
            title = course['title'],
            code = course['code']
        )
        db.session.add(new_course)
        db.session.commit()
        return True, "Added new course."
    except:
        return False, "Error occured while adding course."

def get_course(cid):
    return Course.query.filter_by(id=cid).first().__dict__

def get_courses():
    return [u.__dict__ for u in Course.query.all()]

def enroll_student_in_course(student_id, data):
    student = Student.query.filter_by(id=student_id).first()
    new_enrollment = Enrollment(
        course = data['courses_select'],
        student = student_id,
        ay = student.ay,
        semester = student.semester,
        instructor = uid_to_pk(data['instructors_select'], 'instructor'),
        status = 'pending'
    )
    db.session.add(new_enrollment)
    db.session.commit()
    course_str = str(get_course(data['courses_select'])['code'])
    inst_dict = get_instructor(data['instructors_select'])
    inst_name = f"{inst_dict['l_name'].capitalize()}, {inst_dict['f_name']} {inst_dict['m_name']}."
    return f"Enrolled student in {course_str} under {inst_name}"

def enroll_student_in_program(student_id, data):
    student = Student.query.filter_by(id=student_id).first()
    student.ay = data['academic_year']
    student.semester = data['semester']
    student.program = data['programs_select']
    db.session.commit()
    program_info = get_program(data['programs_select'])
    return f"Enrolled student in {program_info['abbr']} for {data['semester']} Semester Academic Year {data['academic_year']}"

def replace_text_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            inline = paragraph.runs
            for i in range(len(inline)):
                if key in inline[i].text:
                    inline[i].text = inline[i].text.replace(key, str(value))

def replace_text_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)

from docx import Document

def auto_fill(input_path, output_path, replacements):
    try:
        doc = Document(input_path)
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)

        for table in doc.tables:
            replace_text_in_table(table, replacements)

        doc.save(output_path)
        return True
    except:
        return False

def _get_full_name(user):
     u = User.query.filter_by(id=user.user).first()
     return f"{u.l_name.capitalize()} {u.f_name}"

def get_dtr_data(enrollment_id):
    enrollment = Enrollment.query.filter_by(id=enrollment_id).first()
    instructor = Instructor.query.filter_by(id=enrollment.instructor).first()
    data = {
        '[Instructor Name]': _get_full_name(instructor),
        '[Month]': datetime.now().strftime("%B")
    }   
    return data

def get_ssgr_data(enrollment_id):
    enrollment = Enrollment.query.filter_by(id=enrollment_id).first()
    course = Course.query.filter_by(id=enrollment.course).first()
    instructor = Instructor.query.filter_by(id=enrollment.instructor).first()
    student = Student.query.filter_by(id=enrollment.student).first()
    program = Program.query.filter_by(id=student.program).first()
    college = College.query.filter_by(id=program.id).first()
    director = Admin.query.first()
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    data = {
        '[College]': college.abbr,
        '[Semester]': enrollment.semester,
        '[Course Code]': course.code,
        '[Units]': enrollment.units,
        '[Academic Year]': enrollment.ay,
        '[Course Title]': course.title,
        '[Day]': days[date.today().weekday()],
        '[Program]': program.abbr,
        '[Year]': '4',
        '[Time]': f"{datetime.now().time().hour}:{datetime.now().time().minute}",
        '[Student Name]': _get_full_name(student),
        '[Grade]': enrollment.grade,
        '[Units]': enrollment.units,
        '[Instructor Name]': _get_full_name(instructor),
        '[Date]': datetime.now().date(),
        '[Director Name]': _get_full_name(director),
    }
    return data

def get_ter_data(enrollment_id):
    enrollment = Enrollment.query.filter_by(id=enrollment_id).first()
    course = Course.query.filter_by(id=enrollment.course).first()
    instructor = Instructor.query.filter_by(id=enrollment.instructor).first()
    student = Student.query.filter_by(id=enrollment.student).first()
    program = Program.query.filter_by(id=student.program).first()
    data = {
        '[Student Name]': _get_full_name(student),
        '[Program]': program.abbr,
        '[Course Code]': course.code,
        '[Course Title]': course.title,
        '[Units]': enrollment.units,
        '[Grade]': enrollment.grade,
        '[Instructor Name]': _get_full_name(instructor)
    }
    return data

from time import sleep

def delete_file(file_path):
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
            print(f"File {file_path} has been deleted.")
        except Exception as e:
            print(f"An error occurred while trying to delete the file: {e}")
    else:
        print(f"The file {file_path} does not exist.")

def submit_grade(enrollment_id, data):
    try:
        enrollment = Enrollment.query.filter_by(id=enrollment_id).first()
        enrollment.grade = data['grade']
        enrollment.status = 'completed'
        file1name = str(uuid4()) + os.path.splitext(data['file1'].filename)[1]
        file1path = os.path.join(app.config['UPLOAD_FOLDER'], 'grades', file1name)
        data['file1'].save(file1path)
        enrollment.form1 = file1name
        file2name = str(uuid4()) + os.path.splitext(data['file2'].filename)[1]
        file2path = os.path.join(app.config['UPLOAD_FOLDER'], 'grades', file2name)
        data['file2'].save(file2path)
        enrollment.form2 = file2name
        file3name = str(uuid4()) + os.path.splitext(data['file3'].filename)[1]
        file3path = os.path.join(app.config['UPLOAD_FOLDER'], 'grades', file3name)
        data['file3'].save(file3path)
        enrollment.form3 = file3name
        db.session.add(enrollment)
        db.session.commit()
        return True, f"Grades have been submitted."
    except Exception as e:
        return False, f"Error: {e}."

def release_honorarium(enrollment_id):
    try:
        enrollment = Enrollment.query.filter_by(id=enrollment_id).first()
        enrollment.honorarium = 'released'
        enrollment.honorarium_released_date = datetime.now()
        db.session.add(enrollment)
        db.session.commit()
        return True, f"Honorarium has been released."
    except Exception as e:
        return False, f"Problem ocucred while releasing honorarium."
    
def undo_release_honorarium(enrollment_id):
    try:
        enrollment = Enrollment.query.filter_by(id=enrollment_id).first()
        enrollment.honorarium = 'onprocess'
        enrollment.honorarium_released_date = None
        db.session.add(enrollment)
        db.session.commit()
        return True, f"Status for this honorarium has been reverted."
    except Exception as e:
        return False, f"Problem occured while undoing action."

def _remove_sa_instance_state(obj):
    if isinstance(obj, list):
        return [_remove_sa_instance_state(item) for item in obj]
    elif isinstance(obj, dict):
        return {key: _remove_sa_instance_state(value) for key, value in obj.items() if key != '_sa_instance_state'}
    else:
        return obj